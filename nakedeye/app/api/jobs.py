from datetime import datetime, timezone
from io import BytesIO
from typing import Optional
import re
from urllib.parse import urlparse
import uuid

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, Response, UploadFile
import httpx
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.database import get_db, JobApplication

router = APIRouter(prefix="/jobs", tags=["jobs"])

PIPELINE_STATUSES = {
    "applied",
    "screening",
    "interview",
    "offer",
    "rejected",
}

ATS_SIGNATURES = {
    "greenhouse": {"label": "Greenhouse", "tips": "Use exact job-title keywords, answer custom questions directly, keep resume headings standard (Experience, Education, Skills)."},
    "workday": {"label": "Workday", "tips": "Upload a clean DOCX/PDF, mirror the job title verbatim, fill every required field completely."},
    "lever": {"label": "Lever", "tips": "Lead with role-fit keywords in your summary, use concise impact bullets, keep LinkedIn/GitHub links current."},
    "ashby": {"label": "Ashby", "tips": "Match must-have skills exactly, keep dates/titles consistent, write direct screening answers."},
    "smartrecruiters": {"label": "SmartRecruiters", "tips": "Use standard section labels, include certifications prominently, avoid graphics that hide text."},
    "icims": {"label": "iCIMS", "tips": "Prioritize exact job keywords, use traditional resume structure, complete all profile fields."},
    "jobvite": {"label": "Jobvite", "tips": "Keep skills near the top, use measurable bullets, include recruiter-friendly contact details."},
    "successfactors": {"label": "SAP SuccessFactors", "tips": "Use clean resume headings, mirror exact job-post terms, keep profile fields consistent."},
    "taleo": {"label": "Oracle Taleo", "tips": "Use a simple format with no tables/columns, match job title exactly, spell out acronyms."},
    "bamboohr": {"label": "BambooHR", "tips": "Keep formatting minimal, use standard section headers, include all requested info in application fields."},
    "jazzhr": {"label": "JazzHR", "tips": "Mirror job posting language, use keyword-rich bullets, keep resume under 2 pages."},
    "breezy": {"label": "Breezy HR", "tips": "Use clear section headers, include relevant skills early, keep formatting simple."},
    "recruitee": {"label": "Recruitee", "tips": "Match listed requirements point-by-point, use measurable achievements, keep layout clean."},
    "pinpoint": {"label": "Pinpoint", "tips": "Use role-specific keywords, highlight transferable skills, keep resume ATS-friendly."},
    "rippling": {"label": "Rippling", "tips": "Mirror job description language, quantify achievements, use standard resume sections."},
    "workable": {"label": "Workable", "tips": "Include exact skill keywords, use reverse-chronological format, add a targeted summary."},
    "teamtailor": {"label": "Teamtailor", "tips": "Match company culture language, use role-specific keywords, keep formatting clean."},
    "personio": {"label": "Personio", "tips": "Include EU/local qualifications if applicable, mirror job requirements exactly, use standard formatting."},
    "zoho_recruit": {"label": "Zoho Recruit", "tips": "Use standard headers, include all requested skills, keep bullets concise and keyword-rich."},
    "bullhorn": {"label": "Bullhorn", "tips": "Focus on skill tags and certifications, use recruiter-friendly formatting, include contact details prominently."},
    "avature": {"label": "Avature", "tips": "Mirror corporate language, use standard sections, include relevant certifications and clearances."},
    "cornerstone": {"label": "Cornerstone OnDemand", "tips": "Complete all profile sections, use exact job keywords, keep formatting traditional."},
    "phenom": {"label": "Phenom People", "tips": "Use AI-friendly formatting, include skill keywords early, match job title exactly."},
    "adp": {"label": "ADP Workforce", "tips": "Use standard resume format, include all compliance/certification info, mirror job requirements."},
    "paylocity": {"label": "Paylocity", "tips": "Keep formatting simple, use exact keywords from posting, complete all application fields."},
    "paycom": {"label": "Paycom", "tips": "Mirror job description terms, use clean formatting, include all required qualifications."},
    "ukg": {"label": "UKG (Kronos)", "tips": "Use standard sections, include shift/schedule flexibility if relevant, mirror posting keywords."},
    "applytrak": {"label": "ApplyTracker", "tips": "Use keyword-rich bullets, standard formatting, and include all requested information."},
    "myworkdayjobs": {"label": "Workday", "tips": "Upload a clean DOCX/PDF, mirror the job title verbatim, fill every required field completely."},
    "linkedin": {"label": "LinkedIn Jobs", "tips": "Use Easy Apply keywords, match headline to job title, keep profile skills updated."},
    "indeed": {"label": "Indeed", "tips": "Mirror job title and key skills exactly, use Indeed's resume builder format, include location."},
    "hirebridge": {"label": "HireBridge", "tips": "Use traditional formatting, include all listed requirements, keep resume keyword-dense."},
    "clearcompany": {"label": "ClearCompany", "tips": "Match required skills exactly, use clean formatting, include measurable achievements."},
    "jazz": {"label": "JazzHR", "tips": "Mirror job posting language, use keyword-rich bullets, keep resume under 2 pages."},
    "dover": {"label": "Dover", "tips": "Highlight technical skills prominently, use clean formatting, include GitHub/portfolio links."},
    "jobbank": {"label": "Job Bank (Canada)", "tips": "Use NOC codes and exact job title, include Canadian work authorization, highlight bilingual skills if applicable."},
    "usajobs": {"label": "USAJobs", "tips": "Use federal resume format (detailed duties, hours/week), include KSAs, match series/grade requirements exactly."},
    "nhsjobs": {"label": "NHS Jobs (UK)", "tips": "Mirror person specification exactly, include NHS-specific competencies, reference band requirements."},
    "monster": {"label": "Monster", "tips": "Use keyword-rich headline, include location preferences, keep resume in standard format."},
    "ziprecruiter": {"label": "ZipRecruiter", "tips": "Match job title keywords exactly, include salary expectations if asked, keep resume concise."},
    "careerbuilder": {"label": "CareerBuilder", "tips": "Use standard resume format, include searchable keywords, keep contact info current."},
    "glassdoor": {"label": "Glassdoor", "tips": "Research company reviews, mirror job description language, include salary expectations."},
    "angellist": {"label": "Wellfound (AngelList)", "tips": "Highlight startup experience, include tech stack prominently, show side projects."},
    "dice": {"label": "Dice", "tips": "List all technical skills with years of experience, use exact technology names, include clearance level if applicable."},
    "seek": {"label": "SEEK", "tips": "Include Australian/NZ work rights, use local qualifications, mirror job ad keywords exactly."},
    "jobstreet": {"label": "JobStreet", "tips": "Include language proficiency, local certifications, and expected salary in local currency."},
    "jobsireland": {"label": "Jobs Ireland", "tips": "Include Irish work authorization, PPS number availability, mirror job ad language, highlight EU qualifications."},
    "reed": {"label": "Reed", "tips": "Use keyword-rich profile, include salary expectations, keep CV in standard UK format."},
    "totaljobs": {"label": "Totaljobs", "tips": "Mirror job ad keywords, include UK work rights, use standard CV structure."},
    "cwjobs": {"label": "CWJobs", "tips": "List tech stack with experience levels, include contractor/permanent preference, add certifications."},
    "naukri": {"label": "Naukri", "tips": "Include notice period, current CTC and expected CTC, list all technical skills with experience."},
    "bayt": {"label": "Bayt", "tips": "Include visa status, language skills (Arabic is a plus), mirror job requirements."},
    "hays": {"label": "Hays", "tips": "Use recruiter-friendly format, include availability and salary expectations, highlight transferable skills."},
    "roberthalf": {"label": "Robert Half", "tips": "Focus on quantifiable achievements, include contract/permanent preference, list certifications prominently."},
    "efinancialcareers": {"label": "eFinancialCareers", "tips": "Highlight financial certifications (CFA, FRM), include regulatory knowledge, quantify deal sizes."},
    "hired": {"label": "Hired", "tips": "Set salary expectations clearly, highlight tech stack, include role preferences and remote availability."},
    "weworkremotely": {"label": "We Work Remotely", "tips": "Emphasize remote work experience, include timezone availability, highlight async communication skills."},
    "remoteok": {"label": "RemoteOK", "tips": "List remote-friendly skills, include timezone and availability, highlight self-management."},
    "flex": {"label": "FlexJobs", "tips": "Highlight remote/flexible work experience, include schedule preferences, mirror job keywords."},
    "stepstone": {"label": "StepStone", "tips": "Include EU work authorization, language skills, mirror German/EU job ad requirements."},
    "xing": {"label": "XING", "tips": "Optimize DACH-region profile, include German language level, highlight EU qualifications."},
}

ATS_SIGNATURE_PATTERNS = [
    # High-confidence URL patterns
    ("SAP SuccessFactors", 98, ("successfactors.com", "jobs.sap.com", "career4.successfactors", "performancemanager")),
    ("Workday", 98, ("myworkdayjobs.com", "myworkdaysite.com", "wd1.myworkdaysite", "wd3.myworkdaysite", "wd5.myworkdaysite")),
    ("Greenhouse", 98, ("boards.greenhouse.io", "greenhouse.io/", "job_app")),
    ("Lever", 98, ("jobs.lever.co", "lever.co/")),
    ("Ashby", 98, ("jobs.ashbyhq.com", "ashbyhq.com")),
    ("Oracle Taleo", 96, ("taleo.net", "oracle.taleo", "taleoapply", "taleo.com")),
    ("SmartRecruiters", 96, ("smartrecruiters.com", "jobs.smartrecruiters")),
    ("iCIMS", 96, ("icims.com", "careers-", ".icims.")),
    ("Jobvite", 96, ("jobvite.com", "jobs.jobvite")),
    ("BambooHR", 96, ("bamboohr.com/careers", "bamboohr.com/jobs")),
    ("JazzHR", 96, ("applytojob.com", "jazzhr.com")),
    ("Breezy HR", 96, ("breezy.hr",)),
    ("Recruitee", 96, ("recruitee.com", "careers.recruitee")),
    ("Pinpoint", 96, ("pinpointhq.com",)),
    ("Rippling", 96, ("rippling.com/careers", "ats.rippling")),
    ("Workable", 96, ("apply.workable.com", "workable.com/j/")),
    ("Teamtailor", 96, ("teamtailor.com", "career.teamtailor")),
    ("Personio", 96, ("personio.de", "jobs.personio")),
    ("Zoho Recruit", 96, ("zoho.com/recruit", "recruit.zoho")),
    ("Bullhorn", 96, ("bullhornstaffing.com", "bullhorn.com")),
    ("Avature", 96, ("avature.net",)),
    ("Cornerstone OnDemand", 96, ("csod.com", "cornerstone")),
    ("Phenom People", 96, ("phenom.com", "phenompeople")),
    ("ADP Workforce", 94, ("adp.com/careers", "workforcenow.adp")),
    ("Paylocity", 94, ("paylocity.com", "recruiting.paylocity")),
    ("Paycom", 94, ("paycom.com", "paycomonline.net")),
    ("UKG (Kronos)", 94, ("ukg.com", "ultipro.com", "kronos.com")),
    ("LinkedIn Jobs", 90, ("linkedin.com/jobs", "linkedin.com/in/")),
    ("Indeed", 90, ("indeed.com/viewjob", "indeed.com/jobs")),
    ("HireBridge", 94, ("hirebridge.com",)),
    ("ClearCompany", 94, ("clearcompany.com",)),
    ("Dover", 94, ("dover.com", "app.dover.io")),
    # Page-content patterns (lower confidence, detected via inspection)
    ("SAP SuccessFactors", 82, ("rmkcdn.successfactors", "press tab to move to skip to content link", "talent community")),
    ("Workday", 85, ("workday-", "wd-", "workday.com/", "powered by workday")),
    ("Greenhouse", 85, ("greenhouse recruiting", "greenhouse_job_board", "grnh.se")),
    ("Lever", 85, ("lever recruiting", "lever-jobs-embed", "levergreen")),
    ("Oracle Taleo", 82, ("taleo", "requisition", "req id", "job requisition")),
    ("BambooHR", 82, ("bamboohr", "bamboo-hr")),
    ("iCIMS", 82, ("icims", "powered by icims")),
    ("SmartRecruiters", 82, ("smartrecruiters", "smart recruiters")),
    ("Jobvite", 82, ("jobvite", "powered by jobvite")),
    ("Workable", 82, ("powered by workable", "workable-")),
    ("Phenom People", 80, ("phenom", "phenompeople", "talent experience")),
    ("Bullhorn", 80, ("bullhorn", "powered by bullhorn")),
    # Government portals & job boards
    ("Job Bank (Canada)", 98, ("jobbank.gc.ca",)),
    ("USAJobs", 98, ("usajobs.gov",)),
    ("NHS Jobs (UK)", 98, ("jobs.nhs.uk", "trac.jobs")),
    ("Monster", 94, ("monster.com", "jobs.monster")),
    ("ZipRecruiter", 94, ("ziprecruiter.com",)),
    ("CareerBuilder", 94, ("careerbuilder.com",)),
    ("Glassdoor", 94, ("glassdoor.com/job", "glassdoor.com/Job")),
    ("Wellfound (AngelList)", 94, ("wellfound.com", "angel.co/company")),
    ("Dice", 94, ("dice.com",)),
    ("SEEK", 94, ("seek.com.au", "seek.co.nz")),
    ("JobStreet", 94, ("jobstreet.com",)),
    # Page-content fallbacks for government portals
    ("Job Bank (Canada)", 85, ("job bank", "government of canada", "jobbank", "gc.ca")),
    ("USAJobs", 85, ("usajobs", "federal government", "opm.gov")),
    # Regional & international job boards
    ("Jobs Ireland", 98, ("jobsireland.ie",)),
    ("Reed", 96, ("reed.co.uk",)),
    ("Totaljobs", 96, ("totaljobs.com",)),
    ("CWJobs", 96, ("cwjobs.co.uk",)),
    ("Naukri", 96, ("naukri.com",)),
    ("Bayt", 96, ("bayt.com",)),
    ("Hays", 94, ("hays.com", "hays.co.uk", "hays.com.au")),
    ("Robert Half", 94, ("roberthalf.com", "roberthalf.co.uk")),
    ("eFinancialCareers", 94, ("efinancialcareers.com",)),
    ("Hired", 94, ("hired.com",)),
    ("We Work Remotely", 96, ("weworkremotely.com",)),
    ("RemoteOK", 96, ("remoteok.com", "remoteok.io")),
    ("FlexJobs", 94, ("flexjobs.com",)),
    ("StepStone", 96, ("stepstone.de", "stepstone.com")),
    ("XING", 94, ("xing.com",)),
    # Page-content fallbacks for Ireland
    ("Jobs Ireland", 82, ("jobs ireland", "jobsireland", "deasp.gov.ie")),
]

ATS_TIPS = {info["label"]: info["tips"] for info in ATS_SIGNATURES.values()}


STOPWORDS = {
    "about", "above", "across", "after", "again", "against", "also", "and", "any", "are",
    "assigned", "because", "been", "being", "below", "between", "both", "can", "candidate",
    "company", "corporate", "could", "duties", "each", "ensure", "from", "have", "into",
    "must", "other", "our", "perform", "policies", "position", "resources", "relevant",
    "responsibilities", "role", "should", "that", "the", "their", "this", "through",
    "using", "with", "within", "work", "will", "your",
}

LOW_SIGNAL_REQUIREMENT_PATTERNS = (
    "accident free",
    "adhere",
    "assigned",
    "corporate polic",
    "other duties",
    "policies",
    "procedures",
    "satisfy operation",
)

REQUIREMENT_CATEGORIES = {
    "programming_languages": [
        "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust",
        "ruby", "php", "swift", "kotlin", "scala", "perl", "r", "matlab", "lua", "dart",
        "objective-c", "shell", "bash", "powershell", "haskell", "elixir", "clojure",
        "embedded c", "vhdl", "verilog", "assembly", "simulink",
    ],
    "frontend": [
        "react", "reactjs", "react.js", "angular", "vue", "vuejs", "vue.js", "svelte",
        "next.js", "nextjs", "nuxt", "gatsby", "html", "css", "sass", "scss", "less",
        "tailwind", "tailwindcss", "bootstrap", "material ui", "chakra", "webpack", "vite",
        "redux", "zustand", "graphql", "rest api", "responsive design", "accessibility",
        "figma", "ui/ux", "ux design", "ui design", "storybook", "cypress", "playwright",
    ],
    "backend": [
        "node.js", "nodejs", "express", "express.js", "fastapi", "django", "flask",
        "spring boot", "spring", ".net", "asp.net", "rails", "ruby on rails", "laravel",
        "gin", "fiber", "nestjs", "koa", "fastify", "microservices", "rest", "restful",
        "grpc", "websocket", "oauth", "jwt", "authentication", "authorization",
    ],
    "databases": [
        "sql", "mysql", "postgresql", "postgres", "mongodb", "redis", "elasticsearch",
        "dynamodb", "cassandra", "sqlite", "oracle", "sql server", "mariadb", "firebase",
        "firestore", "supabase", "prisma", "sequelize", "typeorm", "mongoose", "nosql",
    ],
    "cloud_devops": [
        "aws", "amazon web services", "azure", "gcp", "google cloud", "docker", "kubernetes",
        "k8s", "terraform", "ansible", "jenkins", "ci/cd", "github actions", "gitlab ci",
        "circleci", "cloudflare", "nginx", "apache", "linux", "ubuntu", "centos",
        "serverless", "lambda", "ec2", "s3", "rds", "ecs", "eks", "fargate", "helm",
        "prometheus", "grafana", "datadog", "new relic", "splunk", "elk stack",
        "infrastructure as code", "iac", "devops", "sre", "site reliability",
        "virtualization", "hypervisor", "vmware", "virtualbox", "container",
    ],
    "data_ai_ml": [
        "machine learning", "deep learning", "artificial intelligence", "ai", "ml",
        "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "scipy",
        "nlp", "natural language processing", "computer vision", "llm", "large language model",
        "gpt", "openai", "langchain", "hugging face", "transformers", "neural network",
        "data science", "data analysis", "data engineering", "etl", "data pipeline",
        "apache spark", "hadoop", "airflow", "dbt", "snowflake", "bigquery", "redshift",
        "tableau", "power bi", "looker", "jupyter", "data visualization",
    ],
    "cybersecurity": [
        "cybersecurity", "information security", "network security", "siem", "soc",
        "firewall", "edr", "ids", "ips", "penetration testing", "pen testing",
        "vulnerability assessment", "threat intelligence", "incident response",
        "encryption", "ssl", "tls", "vpn", "zero trust", "iso 27001", "nist",
        "gdpr", "compliance", "risk assessment", "security audit",
        "owasp", "devsecops", "secure coding", "iec 62443", "nerc cip",
        "anti-tamper", "cybersecurity controls", "risk management framework",
    ],
    "tools_platforms": [
        "git", "github", "gitlab", "bitbucket", "jira", "confluence", "slack",
        "trello", "asana", "notion", "figma", "sketch", "adobe", "photoshop",
        "illustrator", "vs code", "visual studio", "intellij", "postman", "swagger",
        "autocad", "solidworks", "matlab", "labview", "plc", "scada", "hmi",
        "pycharm", "eclipse", "netbeans", "vim", "gdb", "valgrind", "cmake",
        "makefile", "servicenow", "remedy", "hp alm", "rally", "azure devops",
    ],
    "mobile": [
        "ios", "android", "react native", "flutter", "xamarin", "swift", "kotlin",
        "mobile development", "app development", "xcode", "android studio",
        "swiftui", "jetpack compose", "cordova", "ionic", "expo",
    ],
    "testing_qa": [
        "unit testing", "integration testing", "e2e testing", "test automation",
        "selenium", "cypress", "playwright", "jest", "mocha", "pytest", "junit",
        "tdd", "bdd", "qa", "quality assurance", "load testing", "performance testing",
        "robot framework", "robotframework", "regression testing", "smoke testing",
        "test plan", "test case", "manual testing", "firmware testing",
        "automated testing", "test framework", "debugging", "troubleshooting",
        "static analysis", "code coverage", "continuous testing",
    ],
    "soft_skills": [
        "leadership", "communication", "teamwork", "collaboration", "problem solving",
        "critical thinking", "project management", "time management", "presentation",
        "mentoring", "coaching", "stakeholder management", "cross-functional",
        "self-motivated", "detail-oriented", "analytical", "strategic thinking",
    ],
    "methodologies": [
        "agile", "scrum", "kanban", "lean", "waterfall", "sprint", "standup",
        "retrospective", "backlog", "user stories", "continuous improvement",
        "six sigma", "design thinking", "product management", "scaled agile",
        "safe framework", "peer review", "code review", "pair programming",
    ],
    "certifications": [
        "aws certified", "azure certified", "gcp certified", "pmp", "scrum master",
        "csm", "cissp", "cism", "comptia", "security+", "network+", "ccna", "ccnp",
        "ceh", "oscp", "itil", "togaf", "safe", "prince2",
        "bachelor", "master", "degree", "diploma", "phd", "mba",
        "cisco", "microsoft certified", "oracle certified", "osha",
        "certification", "certificate", "license", "licensed", "cysa+",
    ],
    "embedded_systems": [
        "embedded", "firmware", "rtos", "real-time", "microcontroller", "microprocessor",
        "fpga", "arm", "stm32", "arduino", "raspberry pi", "esp32",
        "i2c", "spi", "uart", "can bus", "modbus", "mqtt", "iec 61850",
        "bootloader", "bare metal", "device driver", "interrupt", "dma",
        "system bring-up", "hardware-software", "fault-tolerant", "safety-critical",
        "ardupilot", "px4", "ros", "ros2", "mesh networking", "telemetry",
        "signal processing", "embedded linux", "yocto", "buildroot",
        "jtag", "openocd", "cross-compilation", "toolchain",
    ],
    "networking": [
        "tcp/ip", "dns", "dhcp", "vlan", "vpn", "firewall", "routing", "switching",
        "osi model", "ethernet", "wifi", "bluetooth", "zigbee", "lorawan",
        "network protocol", "socket programming", "http", "https", "ftp", "ssh",
        "snmp", "modbus", "profibus", "profinet", "opc ua", "dnp3",
    ],
    "industry_specific": [
        "automation", "instrumentation", "electrical", "control systems", "embedded",
        "iot", "robotics", "manufacturing", "supply chain", "logistics",
        "healthcare", "fintech", "e-commerce", "saas", "b2b", "b2c",
        "blockchain", "web3", "cryptocurrency", "defi",
        "ar", "vr", "augmented reality", "virtual reality", "gaming", "unreal", "unity",
        "aerospace", "avionics", "defense", "automotive", "energy", "oil and gas",
        "wind power", "solar", "grid", "power systems", "industrial internet",
    ],
    "work_activities": [
        "troubleshooting", "maintaining", "maintenance", "installing", "commissioning",
        "calibration", "repair", "support", "training", "documentation", "reporting",
        "root cause", "continuous improvement", "debugging", "code review",
        "architecture", "system design", "technical writing", "api design",
        "requirements analysis", "system integration", "verification", "validation",
        "risk assessment", "technical documentation", "cross-functional collaboration",
    ],
}

TRACKING_PIXEL_GIF = (
    b"GIF89a\x01\x00\x01\x00\x80\x00\x00\x00\x00\x00"
    b"\xff\xff\xff!\xf9\x04\x01\x00\x00\x00\x00,"
    b"\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;"
)

OPEN_TRACKING_GRACE_SECONDS = 15
OPEN_TRACKING_SCANNER_PATTERNS = (
    "bot",
    "crawler",
    "facebookexternalhit",
    "google-inspectiontool",
    "headless",
    "linkchecker",
    "spider",
    "urlresolver",
    "slurp",
    "wget",
    "curl",
    "python-requests",
    "java/",
    "probe",
)

# Only block known security scanners that auto-load images (not email clients)
EMAIL_PROXY_PATTERNS = (
    "protection.outlook.com",
)


def parse_dt(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    normalized = value.replace("Z", "+00:00")
    parsed = datetime.fromisoformat(normalized)
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed


def normalize_status(status: str) -> str:
    aliases = {
        "ats_screening": "screening",
        "recruiter_review": "screening",
        "assessment": "screening",
        "withdrawn": "rejected",
    }
    return aliases.get(status, status)


def validate_status(status: str) -> str:
    status = normalize_status(status)
    if status not in PIPELINE_STATUSES:
        raise HTTPException(status_code=400, detail="Invalid application status")
    return status


def _ats_result(label: str, confidence: int, evidence: list[str], inspected: bool = False) -> dict:
    return {
        "ats": label,
        "confidence": confidence,
        "tips": ATS_TIPS.get(label, "Use clean formatting, mirror job-post keywords, and avoid image-only resume content."),
        "evidence": evidence[:5],
        "inspected": inspected,
    }


def _detect_ats_signature(text: str, inspected: bool = False) -> Optional[dict]:
    raw = text.lower()
    for label, confidence, patterns in ATS_SIGNATURE_PATTERNS:
        hits = [pattern for pattern in patterns if pattern in raw]
        if hits:
            return _ats_result(label, confidence if inspected else min(confidence, 90), hits, inspected)
    return None


def detect_ats_from_url(url: Optional[str]) -> dict:
    raw = (url or "").strip()
    if not raw:
        return {"ats": "Not detected", "confidence": 0, "tips": "Paste a job URL to detect the ATS.", "evidence": [], "inspected": False}

    parsed = urlparse(raw if "://" in raw else f"https://{raw}")
    host_and_path = f"{parsed.netloc}{parsed.path}"
    detected = _detect_ats_signature(host_and_path)
    if detected:
        return detected

    return {
        "ats": "Unknown ATS",
        "confidence": 15,
        "tips": "URL pattern alone was not enough. Use Inspect Page to check scripts, redirects, and page text.",
        "evidence": ["no known ATS signature in URL"],
        "inspected": False,
    }


async def detect_ats_from_page(url: str) -> dict:
    base = detect_ats_from_url(url)
    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            timeout=12,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"},
        ) as client:
            response = await client.get(url)
        page_text = response.text[:200000]
        final_url = str(response.url)

        # Build comprehensive inspection text from multiple signals
        signals = [final_url, page_text]

        # Extract meta tags
        meta_matches = re.findall(r'<meta[^>]+content=["\']([^"\']+)["\']', page_text, re.IGNORECASE)
        signals.extend(meta_matches)

        # Extract script sources
        script_srcs = re.findall(r'<script[^>]+src=["\']([^"\']+)["\']', page_text, re.IGNORECASE)
        signals.extend(script_srcs)

        # Extract form actions
        form_actions = re.findall(r'<form[^>]+action=["\']([^"\']+)["\']', page_text, re.IGNORECASE)
        signals.extend(form_actions)

        # Extract link hrefs (stylesheets, canonical)
        link_hrefs = re.findall(r'<link[^>]+href=["\']([^"\']+)["\']', page_text, re.IGNORECASE)
        signals.extend(link_hrefs)

        # Extract iframe sources
        iframe_srcs = re.findall(r'<iframe[^>]+src=["\']([^"\']+)["\']', page_text, re.IGNORECASE)
        signals.extend(iframe_srcs)

        inspected_text = " ".join(signals)
        detected = _detect_ats_signature(inspected_text, inspected=True)
        if detected:
            detected["final_url"] = final_url
            return detected
        base["inspected"] = True
        base["final_url"] = final_url
        base["tips"] = "Page was reachable, but no supported ATS signature was found in URL, scripts, meta tags, or page content."
        return base
    except Exception as exc:
        base["warning"] = f"Page inspection failed: {type(exc).__name__}"
        return base


def _tokens(text: Optional[str]) -> list[str]:
    return re.findall(r"[a-z0-9+#./]{2,}", (text or "").lower())


def _normalize_text(text: Optional[str]) -> str:
    return " ".join(_tokens(text))


def _contains_term(text: str, term: str) -> bool:
    pattern = r"(?<![a-z0-9+#./])" + re.escape(term.lower()) + r"(?![a-z0-9+#./])"
    return re.search(pattern, text) is not None


def _is_low_signal(term: str) -> bool:
    lowered = term.lower()
    return any(pattern in lowered for pattern in LOW_SIGNAL_REQUIREMENT_PATTERNS)


# Weight tiers for categories
_CATEGORY_WEIGHTS = {
    "programming_languages": 9,
    "frontend": 8,
    "backend": 8,
    "databases": 8,
    "cloud_devops": 8,
    "data_ai_ml": 8,
    "cybersecurity": 7,
    "tools_platforms": 6,
    "mobile": 8,
    "testing_qa": 6,
    "certifications": 7,
    "industry_specific": 5,
    "methodologies": 4,
    "soft_skills": 3,
    "work_activities": 4,
}

def _detect_jd_section(text: str) -> str:
    """Detect if text is in a 'required' or 'preferred/nice-to-have' section."""
    required_markers = [
        "required", "must have", "must-have", "requirements", "qualifications",
        "what you need", "what we're looking for", "minimum qualifications",
        "basic qualifications", "you must", "essential",
    ]
    preferred_markers = [
        "preferred", "nice to have", "nice-to-have", "bonus", "plus",
        "desired", "preferred qualifications", "additional qualifications",
        "it would be great if", "ideally",
    ]
    lowered = text.lower()
    for marker in preferred_markers:
        if marker in lowered:
            return "preferred"
    for marker in required_markers:
        if marker in lowered:
            return "required"
    return "required"

def _add_requirement(
    requirements: dict[str, dict],
    term: str,
    category: str,
    weight: int,
    source: str,
):
    normalized = " ".join(_tokens(term))
    if not normalized or normalized in STOPWORDS or _is_low_signal(normalized):
        return
    if normalized in requirements:
        requirements[normalized]["weight"] = max(requirements[normalized]["weight"], weight)
        if source not in requirements[normalized]["sources"]:
            requirements[normalized]["sources"].append(source)
        return
    requirements[normalized] = {
        "term": normalized,
        "category": category,
        "weight": weight,
        "sources": [source],
    }

def _extract_requirements(job_description: Optional[str]) -> list[dict]:
    jd_text = _normalize_text(job_description)
    jd_raw = (job_description or "").lower()
    requirements: dict[str, dict] = {}

    # Detect if JD has preferred/required sections
    section_type = _detect_jd_section(jd_raw)
    is_preferred_section = section_type == "preferred"

    # Scan known skill categories against job description
    for category, terms in REQUIREMENT_CATEGORIES.items():
        base_weight = _CATEGORY_WEIGHTS.get(category, 5)
        for term in terms:
            if _contains_term(jd_text, term):
                weight = max(2, base_weight - 3) if is_preferred_section else base_weight
                _add_requirement(requirements, term, category, weight, "known-skill")

    # Extract years-of-experience mentions
    exp_patterns = re.findall(r"(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)", jd_raw)
    for years in exp_patterns:
        _add_requirement(requirements, f"{years}+ years experience", "certifications", 6, "experience-req")

    # Extract degree requirements
    degree_patterns = re.findall(r"(bachelor'?s?|master'?s?|ph\.?d|mba|associate'?s?)\s+(?:degree|in\s)", jd_raw)
    for deg in degree_patterns:
        _add_requirement(requirements, "degree", "certifications", 6, "education-req")
        _add_requirement(requirements, deg.strip(), "certifications", 5, "education-req")

    return sorted(requirements.values(), key=lambda item: (-item["weight"], item["term"]))[:40]

def analyze_strength(job_description: Optional[str], resume_keywords: Optional[str]) -> dict:
    requirements = _extract_requirements(job_description)
    resume_text = _normalize_text(resume_keywords)
    resume_terms = set(_tokens(resume_keywords))

    matched, missing, quick_wins = [], [], []
    category_summary: dict[str, dict] = {}
    possible_weight = sum(item["weight"] for item in requirements)
    matched_weight = 0
    critical_missing = 0

    for item in requirements:
        term = item["term"]
        weight = item["weight"]
        category = item["category"]
        parts = term.split()
        is_match = _contains_term(resume_text, term) or all(part in resume_terms for part in parts)
        summary = category_summary.setdefault(category, {"matched": [], "missing": []})
        if is_match:
            matched.append(term)
            matched_weight += weight
            summary["matched"].append(term)
        else:
            missing.append(term)
            summary["missing"].append(term)
            if weight >= 7:
                critical_missing += 1
            # Quick wins: easy-to-add missing keywords
            if weight <= 5 and len(term.split()) <= 2:
                quick_wins.append(term)

    raw_score = round((matched_weight / possible_weight) * 100) if possible_weight else 0
    evidence_terms = len(resume_terms)
    evidence_cap = min(100, max(25, evidence_terms * 5))
    score = min(raw_score, evidence_cap)

    # Apply penalty for missing critical skills
    if critical_missing >= 3:
        score = max(0, score - (critical_missing * 3))

    # Determine verdict
    if score >= 80:
        verdict = "Strong Match"
        verdict_detail = "Your resume aligns well with this role. Focus on tailoring your summary and quantifying achievements."
    elif score >= 55:
        verdict = "Moderate Match"
        verdict_detail = "You have a decent foundation. Add missing technical keywords and highlight relevant experience."
    elif score >= 30:
        verdict = "Weak Match"
        verdict_detail = "Significant gaps exist. Consider adding missing skills to your resume or gaining relevant experience."
    else:
        verdict = "Low Match"
        verdict_detail = "Major skill gaps detected. This role may require significant upskilling or may not be the right fit."

    warnings = []
    if len(requirements) < 5:
        warnings.append("Not enough high-signal requirements found. Paste more of the job description for better accuracy.")
    if evidence_terms < 5:
        warnings.append("Resume evidence is too thin; upload the full CV or paste more resume text.")
    if score != raw_score and evidence_terms < 20:
        warnings.append("Score capped due to limited resume evidence.")

    return {
        "score": score,
        "raw_score": raw_score,
        "verdict": verdict,
        "verdict_detail": verdict_detail,
        "matched_keywords": matched[:15],
        "missing_keywords": missing[:15],
        "quick_wins": quick_wins[:8],
        "target_keywords": [item["term"] for item in requirements[:20]],
        "category_summary": category_summary,
        "critical_missing": critical_missing,
        "total_requirements": len(requirements),
        "warnings": warnings,
    }



def _clean_extracted_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs)


def extract_resume_text(filename: str, content: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        text = _extract_pdf_text(content)
    elif lower_name.endswith(".docx"):
        text = _extract_docx_text(content)
    elif lower_name.endswith(".txt"):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail="Upload a PDF, DOCX, or TXT resume.")

    cleaned = _clean_extracted_text(text)
    if len(cleaned) < 80:
        raise HTTPException(status_code=400, detail="Could not extract enough readable text from this resume.")
    return cleaned


def _aware_utc(value: Optional[datetime]) -> Optional[datetime]:
    if value is None:
        return None
    if value.tzinfo is None:
        return value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc)


def _should_record_email_open(request: Request, application: JobApplication, now: datetime) -> bool:
    # Only track if email was actually sent or already opened (allow multiple opens)
    if application.email_tracking_status not in ("sent", "opened"):
        return False

    sent_at = _aware_utc(application.email_sent_at) or _aware_utc(application.updated_at)
    if sent_at and (now - sent_at).total_seconds() < OPEN_TRACKING_GRACE_SECONDS:
        return False

    user_agent = request.headers.get("user-agent", "").lower()
    purpose = " ".join(
        request.headers.get(name, "").lower()
        for name in ("purpose", "sec-purpose", "x-purpose")
    )
    if "prefetch" in purpose:
        return False

    # Filter out known bots and scanners
    if any(pattern in user_agent for pattern in OPEN_TRACKING_SCANNER_PATTERNS):
        return False

    # Filter email proxy auto-loads by checking referrer/via headers
    via = request.headers.get("via", "").lower()
    x_forwarded = request.headers.get("x-forwarded-for", "").lower()
    if any(proxy in via or proxy in x_forwarded for proxy in EMAIL_PROXY_PATTERNS):
        return False

    return True


class JobApplicationCreate(BaseModel):
    company: str
    role: str
    status: str = "applied"
    job_url: Optional[str] = None
    ats_url: Optional[str] = None
    ats_detected: Optional[str] = None
    ats_score: Optional[int] = None
    source: Optional[str] = None
    location: Optional[str] = None
    resume_version: Optional[str] = None
    job_description: Optional[str] = None
    strength_score: Optional[int] = None
    missing_keywords: Optional[str] = None
    contact_name: Optional[str] = None
    contact_email: Optional[str] = None
    email_tracking_status: Optional[str] = None
    tracking_pixel_id: Optional[str] = None
    linkedin_profile_viewed: bool = False
    notes: Optional[str] = None
    applied_at: Optional[str] = None
    follow_up_at: Optional[str] = None
    interview_at: Optional[str] = None


class JobApplicationUpdate(BaseModel):
    company: Optional[str] = None
    role: Optional[str] = None
    status: Optional[str] = None
    job_url: Optional[str] = None
    ats_url: Optional[str] = None
    ats_detected: Optional[str] = None
    ats_score: Optional[int] = None
    source: Optional[str] = None
    location: Optional[str] = None
    resume_version: Optional[str] = None
    job_description: Optional[str] = None
    strength_score: Optional[int] = None
    missing_keywords: Optional[str] = None
    contact_name: Optional[str] = None
    contact_email: Optional[str] = None
    email_tracking_status: Optional[str] = None
    tracking_pixel_id: Optional[str] = None
    linkedin_profile_viewed: Optional[bool] = None
    notes: Optional[str] = None
    applied_at: Optional[str] = None
    follow_up_at: Optional[str] = None
    interview_at: Optional[str] = None


class AtsDetectRequest(BaseModel):
    job_url: str


class StrengthAnalyzeRequest(BaseModel):
    job_description: str
    resume_keywords: str


class FollowUpEmailRequest(BaseModel):
    application_id: Optional[str] = None
    company: Optional[str] = None
    role: Optional[str] = None
    recipient_name: Optional[str] = None
    dashboard_url: Optional[str] = None


@router.get("/applications")
async def list_applications(status: Optional[str] = None, db: AsyncSession = Depends(get_db)):
    query = select(JobApplication).order_by(JobApplication.applied_at.desc())
    if status:
        query = query.where(JobApplication.status == validate_status(status))
    return (await db.scalars(query)).all()


@router.get("/stats")
async def get_application_stats(db: AsyncSession = Depends(get_db)):
    applications = (await db.scalars(select(JobApplication))).all()
    total = len(applications)
    by_status = {status: 0 for status in PIPELINE_STATUSES}
    by_resume = {}
    for app in applications:
        status = normalize_status(app.status)
        by_status[status] = by_status.get(status, 0) + 1
        if app.resume_version:
            by_resume.setdefault(app.resume_version, {"total": 0, "interviews": 0, "offers": 0})
            by_resume[app.resume_version]["total"] += 1
            if status in {"interview", "offer"}:
                by_resume[app.resume_version]["interviews"] += 1
            if status == "offer":
                by_resume[app.resume_version]["offers"] += 1

    response_count = sum(by_status.get(s, 0) for s in ("screening", "interview", "offer", "rejected"))
    interview_count = by_status.get("interview", 0) + by_status.get("offer", 0)
    offer_count = by_status.get("offer", 0)

    return {
        "total": total,
        "by_status": by_status,
        "response_rate": round((response_count / total) * 100, 1) if total else 0,
        "interview_rate": round((interview_count / total) * 100, 1) if total else 0,
        "offer_rate": round((offer_count / total) * 100, 1) if total else 0,
        "by_resume": by_resume,
    }


@router.post("/applications", status_code=201)
async def create_application(data: JobApplicationCreate, db: AsyncSession = Depends(get_db)):
    status = validate_status(data.status)
    application = JobApplication(
        id=str(uuid.uuid4()),
        company=data.company,
        role=data.role,
        status=status,
        job_url=data.job_url,
        ats_url=data.ats_url,
        ats_detected=data.ats_detected or detect_ats_from_url(data.job_url or data.ats_url).get("ats"),
        ats_score=data.ats_score,
        source=data.source,
        location=data.location,
        resume_version=data.resume_version,
        job_description=data.job_description,
        strength_score=data.strength_score,
        missing_keywords=data.missing_keywords,
        contact_name=data.contact_name,
        contact_email=data.contact_email,
        email_tracking_status=data.email_tracking_status or "not_sent",
        tracking_pixel_id=data.tracking_pixel_id or str(uuid.uuid4()),
        linkedin_profile_viewed=data.linkedin_profile_viewed,
        notes=data.notes,
        applied_at=parse_dt(data.applied_at) or datetime.now(timezone.utc),
        follow_up_at=parse_dt(data.follow_up_at),
        interview_at=parse_dt(data.interview_at),
    )
    db.add(application)
    await db.commit()
    await db.refresh(application)
    return application


@router.patch("/applications/{application_id}")
async def update_application(application_id: str, data: JobApplicationUpdate, db: AsyncSession = Depends(get_db)):
    application = await db.get(JobApplication, application_id)
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    payload = data.model_dump(exclude_none=True)
    for field, value in payload.items():
        if field == "status":
            value = validate_status(value)
        elif field in {"applied_at", "follow_up_at", "interview_at"}:
            value = parse_dt(value)
        setattr(application, field, value)
    application.updated_at = datetime.now(timezone.utc)
    await db.commit()
    await db.refresh(application)
    return application


@router.post("/ats-detect")
async def detect_ats(data: AtsDetectRequest):
    return await detect_ats_from_page(data.job_url)


@router.post("/strength-analyze")
async def strength_analyze(data: StrengthAnalyzeRequest):
    return analyze_strength(data.job_description, data.resume_keywords)


@router.post("/resume-extract")
async def resume_extract(resume: UploadFile = File(...)):
    content = await resume.read()
    if len(content) > 6 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Resume file is too large. Use a file under 6 MB.")
    text = extract_resume_text(resume.filename or "", content)
    return {
        "filename": resume.filename,
        "characters": len(text),
        "preview": text[:500],
        "text": text,
    }


@router.post("/follow-up-email")
async def generate_follow_up_email(data: FollowUpEmailRequest, db: AsyncSession = Depends(get_db)):
    application = None
    if data.application_id:
        application = await db.get(JobApplication, data.application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")

    company = data.company or (application.company if application else "the team")
    role = data.role or (application.role if application else "the role")
    recipient = data.recipient_name or (application.contact_name if application else None) or "there"
    email_body = (
        f"Hi {recipient},\n\n"
        f"I hope you are doing well. I wanted to follow up on my application for the {role} position at {company}. "
        "I remain very interested in the opportunity and would appreciate any update you can share on the next steps.\n\n"
        "Thank you for your time,\n"
        "[Your Name]"
    )

    pixel_id = (application.tracking_pixel_id if application else None) or str(uuid.uuid4())

    return {"subject": f"Following up on {role} application", "body": email_body, "tracking_pixel_id": pixel_id}


@router.get("/email-open/{pixel_id}.gif")
async def track_email_open(pixel_id: str, request: Request, db: AsyncSession = Depends(get_db)):
    import logging
    logger = logging.getLogger("nakedeye.email-tracking")
    ua = request.headers.get("user-agent", "unknown")
    logger.info(f"[PIXEL HIT] pixel_id={pixel_id} ua={ua[:80]} ip={request.client.host if request.client else '?'}")
    application = await db.scalar(select(JobApplication).where(JobApplication.tracking_pixel_id == pixel_id))
    now = datetime.now(timezone.utc)
    if application:
        should_record = _should_record_email_open(request, application, now)
        logger.info(f"[PIXEL] app={application.company}/{application.role} status={application.email_tracking_status} should_record={should_record}")
        if should_record:
            application.email_tracking_status = "opened"
            application.email_opened_at = now
            application.email_open_count = (application.email_open_count or 0) + 1
            application.updated_at = now
            await db.commit()
            logger.info(f"[PIXEL] ✅ RECORDED OPEN #{application.email_open_count} for {application.company}")
    else:
        logger.warning(f"[PIXEL] ❌ No application found for pixel_id={pixel_id}")
    return Response(content=TRACKING_PIXEL_GIF, media_type="image/gif", headers={"Cache-Control": "no-store, no-cache, must-revalidate", "Pragma": "no-cache"})




class SendFollowUpRequest(BaseModel):
    application_id: str
    subject: str
    body: str
    dashboard_url: Optional[str] = None


@router.post("/send-follow-up")
async def send_follow_up_email(data: SendFollowUpRequest, db: AsyncSession = Depends(get_db)):
    """
    Sends the follow-up email directly via SMTP using the configured credentials.
    Appends the tracking pixel invisibly so the recruiter never sees it.
    """
    import html
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from app.config import settings

    application = await db.get(JobApplication, data.application_id)
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    if not settings.SMTP_USER or not settings.SMTP_PASS:
        raise HTTPException(status_code=400, detail="SMTP not configured — set SMTP_USER and SMTP_PASS in Notifications settings")

    if not application.contact_email:
        raise HTTPException(status_code=400, detail="Add Contact Email to this application before sending a follow-up")

    pixel_id = str(uuid.uuid4())
    # Use ngrok public URL for tracking pixel so recruiters can reach it
    base_url = (data.dashboard_url or "").strip().rstrip("/")
    # If running behind ngrok, use the ngrok domain
    if "localhost" in base_url or "127.0.0.1" in base_url or not base_url:
        from app.config import settings
        ngrok_url = getattr(settings, 'NGROK_URL', '') or 'https://reissue-lusty-tattle.ngrok-free.dev'
        base_url = ngrok_url.rstrip('/')
    pixel_url = f"{base_url}/api/jobs/email-open/{pixel_id}.gif?ngrok-skip-browser-warning=true"

    # Convert plain text body to HTML — preserves line breaks, appends invisible pixel
    safe_body = html.escape(data.body).replace("\n", "<br>")
    html_body = (
        "<html><body style='font-family:Arial,sans-serif;font-size:14px;color:#222;line-height:1.6;'>"
        + safe_body
        + f'<img src="{pixel_url}" width="1" height="1" alt="" style="display:none;"/>'
        + "</body></html>"
    )

    msg = MIMEMultipart("alternative")
    msg["Subject"] = data.subject
    msg["From"]    = f"Alexander Sugian <{settings.SMTP_USER}>"
    msg["To"]      = application.contact_email
    msg.attach(MIMEText(data.body, "plain"))
    msg.attach(MIMEText(html_body, "html"))

    try:
        with smtplib.SMTP(settings.SMTP_HOST or "smtp.gmail.com", settings.SMTP_PORT or 587, timeout=20) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(settings.SMTP_USER, settings.SMTP_PASS)
            server.sendmail(settings.SMTP_USER, [application.contact_email], msg.as_string())

        # Update tracking status
        now = datetime.now(timezone.utc)
        application.email_tracking_status = "sent"
        application.tracking_pixel_id = pixel_id
        application.email_sent_at = now
        application.email_opened_at = None
        application.email_open_count = 0
        application.updated_at = now
        await db.commit()
        return {"success": True, "recipient": msg["To"], "pixel_url": pixel_url}
    except smtplib.SMTPAuthenticationError:
        return {"success": False, "error": "SMTP authentication failed. Check Gmail App Password in Notifications."}
    except smtplib.SMTPRecipientsRefused:
        return {"success": False, "error": "Recipient was refused. Check the application's Contact Email."}
    except smtplib.SMTPConnectError:
        return {"success": False, "error": "Could not connect to SMTP server. Check SMTP host/port or network access."}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/application-email")
async def generate_application_email(data: FollowUpEmailRequest, db: AsyncSession = Depends(get_db)):
    """Generate a professional application email draft."""
    application = None
    if data.application_id:
        application = await db.get(JobApplication, data.application_id)
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")

    company = data.company or (application.company if application else "the team")
    role = data.role or (application.role if application else "the role")
    recipient = data.recipient_name or (application.contact_name if application else None) or "Hiring Manager"
    email_body = (
        f"Dear {recipient},\n\n"
        f"I am writing to express my interest in the {role} position at {company}. "
        "Please find my resume attached for your review.\n\n"
        f"I am confident that my skills and experience make me a strong candidate for this role, "
        "and I would welcome the opportunity to discuss how I can contribute to your team.\n\n"
        "Thank you for considering my application. I look forward to hearing from you.\n\n"
        "Best regards,\n"
        "[Your Name]"
    )

    return {"subject": f"Application for {role} — {company}", "body": email_body}


@router.post("/send-application")
async def send_application_email(
    application_id: str = Form(""),
    subject: str = Form(""),
    body: str = Form(""),
    dashboard_url: str = Form(""),
    resume: UploadFile = File(None),
    db: AsyncSession = Depends(get_db),
):
    """Send application email with resume attached and tracking pixel embedded."""
    import html as html_module
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    from app.config import settings

    if not application_id:
        raise HTTPException(status_code=400, detail="Application ID is required")

    application = await db.get(JobApplication, application_id)
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    if not settings.SMTP_USER or not settings.SMTP_PASS:
        raise HTTPException(status_code=400, detail="SMTP not configured — set SMTP_USER and SMTP_PASS in Notifications settings")

    if not application.contact_email:
        raise HTTPException(status_code=400, detail="Add Contact Email to this application before sending")

    if not subject or not body:
        raise HTTPException(status_code=400, detail="Subject and message body are required")

    pixel_id = str(uuid.uuid4())
    base_url = (dashboard_url or "").strip().rstrip("/")
    if "localhost" in base_url or "127.0.0.1" in base_url or not base_url:
        ngrok_url = getattr(settings, 'NGROK_URL', '') or 'https://reissue-lusty-tattle.ngrok-free.dev'
        base_url = ngrok_url.rstrip('/')
    pixel_url = f"{base_url}/api/jobs/email-open/{pixel_id}.gif?ngrok-skip-browser-warning=true"

    safe_body = html_module.escape(body).replace("\n", "<br>")
    html_body = (
        "<html><body style='font-family:Arial,sans-serif;font-size:14px;color:#222;line-height:1.6;'>"
        + safe_body
        + f'<img src="{pixel_url}" width="1" height="1" alt="" style="display:none;"/>'
        + "</body></html>"
    )

    msg = MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"] = f"Alexander Sugian <{settings.SMTP_USER}>"
    msg["To"] = application.contact_email

    # Store resume and create tracked download link
    resume_download_id = str(uuid.uuid4())
    resume_download_url = None
    attached_name = None
    if resume and resume.filename:
        file_content = await resume.read()
        if len(file_content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Resume file too large. Max 10 MB.")

        # Save resume to disk for tracked downloads
        import os
        resume_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "resumes")
        os.makedirs(resume_dir, exist_ok=True)
        ext = os.path.splitext(resume.filename)[1] or ".pdf"
        resume_path = os.path.join(resume_dir, f"{resume_download_id}{ext}")
        with open(resume_path, "wb") as f:
            f.write(file_content)

        attached_name = resume.filename
        resume_download_url = f"{base_url}/api/jobs/resume-download/{resume_download_id}{ext}?ngrok-skip-browser-warning=true"

        # Also attach to email directly
        attachment = MIMEBase("application", "octet-stream")
        attachment.set_payload(file_content)
        encoders.encode_base64(attachment)
        attachment.add_header("Content-Disposition", f"attachment; filename={resume.filename}")
        msg.attach(attachment)

    html_body = (
        "<html><body style='font-family:Arial,sans-serif;font-size:14px;color:#222;line-height:1.6;'>"
        + safe_body
        + f'<img src="{pixel_url}" width="1" height="1" alt="" style="display:none;"/>'
        + "</body></html>"
    )

    body_part = MIMEMultipart("alternative")
    body_part.attach(MIMEText(body, "plain"))
    body_part.attach(MIMEText(html_body, "html"))
    msg.attach(body_part)

    try:
        with smtplib.SMTP(settings.SMTP_HOST or "smtp.gmail.com", settings.SMTP_PORT or 587, timeout=20) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(settings.SMTP_USER, settings.SMTP_PASS)
            server.sendmail(settings.SMTP_USER, [application.contact_email], msg.as_string())

        now = datetime.now(timezone.utc)
        application.email_tracking_status = "sent"
        application.tracking_pixel_id = pixel_id
        application.email_sent_at = now
        application.email_opened_at = None
        application.email_open_count = 0
        if resume and resume.filename:
            application.resume_download_id = resume_download_id
            application.resume_filename = attached_name
            application.resume_downloaded_at = None
            application.resume_download_count = 0
        application.updated_at = now
        await db.commit()
        return {"success": True, "recipient": msg["To"], "pixel_url": pixel_url, "attached": attached_name, "resume_tracked": bool(resume_download_url)}
    except smtplib.SMTPAuthenticationError:
        return {"success": False, "error": "SMTP authentication failed. Check Gmail App Password in Notifications."}
    except smtplib.SMTPRecipientsRefused:
        return {"success": False, "error": "Recipient was refused. Check the application's Contact Email."}
    except smtplib.SMTPConnectError:
        return {"success": False, "error": "Could not connect to SMTP server. Check SMTP host/port or network access."}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.get("/resume-download/{download_id}")
async def track_resume_download(download_id: str, request: Request, db: AsyncSession = Depends(get_db)):
    """Serve the resume file and record that it was downloaded."""
    import os

    # Strip extension from download_id for DB lookup
    base_id = os.path.splitext(download_id)[0]

    application = await db.scalar(
        select(JobApplication).where(JobApplication.resume_download_id == base_id)
    )

    # Find the file on disk
    resume_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "resumes")
    matching = [f for f in os.listdir(resume_dir) if f.startswith(base_id)] if os.path.isdir(resume_dir) else []
    if not matching:
        raise HTTPException(status_code=404, detail="Resume not found")

    filepath = os.path.join(resume_dir, matching[0])
    filename = application.resume_filename if application else matching[0]

    # Record the download (filter out bots)
    user_agent = request.headers.get("user-agent", "").lower()
    is_bot = any(p in user_agent for p in OPEN_TRACKING_SCANNER_PATTERNS)
    if application and not is_bot:
        now = datetime.now(timezone.utc)
        application.resume_downloaded_at = now
        application.resume_download_count = (application.resume_download_count or 0) + 1
        application.updated_at = now
        await db.commit()

    # Determine content type
    ext = os.path.splitext(filepath)[1].lower()
    content_types = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
    }
    content_type = content_types.get(ext, "application/octet-stream")

    with open(filepath, "rb") as f:
        content = f.read()

    return Response(
        content=content,
        media_type=content_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.delete("/applications/{application_id}", status_code=204)
async def delete_application(application_id: str, db: AsyncSession = Depends(get_db)):
    application = await db.get(JobApplication, application_id)
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    await db.delete(application)
    await db.commit()

